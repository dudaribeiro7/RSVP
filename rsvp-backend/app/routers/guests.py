# app/routers/guests.py
from __future__ import annotations

from datetime import datetime, timezone

from typing import List

from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session

from app.database import get_db
from app import models, schemas

from app.security import require_admin

from io import BytesIO
from datetime import datetime

from fastapi.responses import StreamingResponse

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RLTable, TableStyle
from reportlab.lib.styles import getSampleStyleSheet


router = APIRouter(
    prefix="/guests",
    tags=["Guests"],
)

def ensure_utc(dt):
    if not dt:
        return dt
    if dt.tzinfo is None:
        return dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


# ==========================
#  Normalização de nomes
# ==========================
def normalize_name(name: str) -> str:
    """
    Normaliza nomes deixando cada palavra capitalizada.
    Ex: 'mARIA eduARDA fACIO' -> 'Maria Eduarda Facio'
    """
    if not name:
        return name

    return " ".join(word.capitalize() for word in name.split())


# ==========================
#  CREATE GUEST (RSVP)
# ==========================
@router.post("/", response_model=schemas.GuestResponse, status_code=status.HTTP_201_CREATED)
def create_guest(guest: schemas.GuestCreate, db: Session = Depends(get_db)):
    normalized_name = normalize_name(guest.name)

    db_guest = models.Guest(
        name=normalized_name,
        phone=guest.phone,
        rsvp_status=guest.rsvp_status.value if hasattr(guest.rsvp_status, "value") else str(guest.rsvp_status),
        note=guest.note,
        responded_at=datetime.now(timezone.utc),
    )
    db.add(db_guest)
    db.commit()
    db.refresh(db_guest)

    # Só cria acompanhantes se a pessoa marcou que VAI.
    if db_guest.rsvp_status == schemas.RSVPStatus.YES.value:
        for comp in guest.companions:
            normalized_comp_name = normalize_name(comp.name)
            new_comp = models.Companion(
                name=normalized_comp_name,
                guest_id=db_guest.id,
            )
            db.add(new_comp)

        db.commit()
        db.refresh(db_guest)

    return db_guest


# ==========================
#  LIST ALL GUESTS
# ==========================
@router.get("/", response_model=List[schemas.GuestResponse])
def list_guests(
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    guests = db.query(models.Guest).all()
    for g in guests:
        g.responded_at = ensure_utc(g.responded_at)
    return guests




# ==========================
#  FIND GUEST BY q
#  (ID, name, phone)
# ==========================
@router.get("/find", response_model=List[schemas.GuestResponse])
def find_guests(
    q: str,
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    possible_id = int(q) if q.isdigit() else None

    guests = (
        db.query(models.Guest)
        .filter(
            (models.Guest.id == possible_id)
            | (models.Guest.name.ilike(f"%{q}%"))
            | (models.Guest.phone.ilike(f"%{q}%"))
        )
        .all()
    )

    return guests


# ==========================
#  GET GUEST BY ID
# ==========================
@router.get("/{guest_id}", response_model=schemas.GuestResponse)
def get_guest(
    guest_id: int,
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    guest = db.query(models.Guest).filter(models.Guest.id == guest_id).first()
    if not guest:
        raise HTTPException(404, "Convidado não encontrado.")
    return guest


# ==========================
#  UPDATE GUEST (PATCH)
#  (útil se você criar um admin futuramente)
# ==========================
@router.patch("/{guest_id}", response_model=schemas.GuestResponse)
def update_guest(
    guest_id: int,
    data: schemas.GuestUpdate,
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    guest = db.query(models.Guest).filter(models.Guest.id == guest_id).first()
    if not guest:
        raise HTTPException(404, "Convidado não encontrado.")

    if data.name is not None:
        guest.name = normalize_name(data.name)

    if data.phone is not None:
        guest.phone = data.phone

    # Se alterar status/recado, atualiza responded_at
    status_changed = False

    if data.rsvp_status is not None:
        new_status = data.rsvp_status.value if hasattr(data.rsvp_status, "value") else str(data.rsvp_status)
        if new_status != guest.rsvp_status:
            guest.rsvp_status = new_status
            status_changed = True

    if data.note is not None:
        guest.note = data.note
        status_changed = True

    # Se mudou pra NO/MAYBE, remove acompanhantes (não faz sentido manter)
    if status_changed:
        guest.responded_at = datetime.now(timezone.utc)

        if guest.rsvp_status in {schemas.RSVPStatus.NO.value, schemas.RSVPStatus.MAYBE.value}:
            guest.companions.clear()  # cascade delete-orphan

    db.commit()
    db.refresh(guest)
    return guest


# ==========================
#  DELETE GUEST
# ==========================
@router.delete("/{guest_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_guest(
    guest_id: int,
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    guest = db.query(models.Guest).filter(models.Guest.id == guest_id).first()
    if not guest:
        raise HTTPException(404, "Convidado não encontrado.")
    db.delete(guest)
    db.commit()
    return


def _confirmed_attendees_with_tables(db: Session) -> list[str]:
    """Retorna lista de strings 'Nome - Mesa X' (ou só 'Nome' se sem mesa),
    ordenada alfabeticamente."""
    confirmed_guests = (
        db.query(models.Guest)
        .filter(models.Guest.rsvp_status == schemas.RSVPStatus.YES.value)
        .order_by(models.Guest.name.asc())
        .all()
    )

    # Monta mapa person_id -> table_number a partir de TableArrangement
    arrangements = db.query(models.TableArrangement).all()
    guest_table: dict[int, int] = {}
    companion_table: dict[int, int] = {}
    for a in arrangements:
        if a.guest_id is not None:
            guest_table[a.guest_id] = a.table_number
        if a.companion_id is not None:
            companion_table[a.companion_id] = a.table_number

    entries: list[str] = []
    for g in confirmed_guests:
        if g.name:
            tbl = guest_table.get(g.id)
            label = f"{g.name.strip()} - Mesa {tbl}" if tbl else g.name.strip()
            entries.append(label)

        for c in (g.companions or []):
            if c.name:
                tbl = companion_table.get(c.id)
                label = f"{c.name.strip()} - Mesa {tbl}" if tbl else c.name.strip()
                entries.append(label)

    # remove duplicados e ordena por casefold
    unique = sorted(set(entries), key=lambda s: s.casefold())
    return unique


@router.get("/export/confirmed.docx")
def export_confirmed_docx(
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    names = _confirmed_attendees_with_tables(db)

    doc = Document()
    doc.add_heading("Convidados confirmados", level=1)
    doc.add_paragraph(f"Total: {len(names)}")
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    # ---------- tabela invisível de 2 colunas ----------
    half = (len(names) + 1) // 2          # coluna esquerda pega o item extra
    col_left = names[:half]
    col_right = names[half:]

    table = doc.add_table(rows=max(len(col_left), len(col_right)), cols=2)
    table.autofit = True

    # remove bordas da tabela (visual limpo)
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0",
            qn("w:space"): "0", qn("w:color"): "auto",
        })
        borders.append(el)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)

    for i, row in enumerate(table.rows):
        left_text = f"{i + 1}. {col_left[i]}" if i < len(col_left) else ""
        right_text = f"{half + i + 1}. {col_right[i]}" if i < len(col_right) else ""
        row.cells[0].text = left_text
        row.cells[1].text = right_text

        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"confirmados-{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/export/confirmed.pdf")
def export_confirmed_pdf(
    db: Session = Depends(get_db),
    admin: None = Depends(require_admin),
):
    names = _confirmed_attendees_with_tables(db)

    bio = BytesIO()
    styles = getSampleStyleSheet()

    pdf = SimpleDocTemplate(bio, pagesize=A4, title="Convidados confirmados")
    story = [
        Paragraph("Convidados confirmados", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Total: {len(names)}", styles["Normal"]),
        Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]),
        Spacer(1, 16),
    ]

    # ---------- tabela invisível de 2 colunas ----------
    half = (len(names) + 1) // 2
    col_left = names[:half]
    col_right = names[half:]

    style_normal = styles["Normal"]
    table_data = []
    for i in range(max(len(col_left), len(col_right))):
        left = Paragraph(f"{i + 1}. {col_left[i]}", style_normal) if i < len(col_left) else ""
        right = Paragraph(f"{half + i + 1}. {col_right[i]}", style_normal) if i < len(col_right) else ""
        table_data.append([left, right])

    page_w = A4[0] - 2 * 72  # largura útil (margens padrão 1 in = 72pt)
    col_w = page_w / 2

    t = RLTable(table_data, colWidths=[col_w, col_w])
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        # sem bordas
        ("LINEBELOW", (0, 0), (-1, -1), 0, "white"),
        ("LINEABOVE", (0, 0), (-1, -1), 0, "white"),
    ]))
    story.append(t)

    pdf.build(story)

    bio.seek(0)
    filename = f"confirmados-{datetime.now().strftime('%Y-%m-%d_%H-%M')}.pdf"
    return StreamingResponse(
        bio,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
